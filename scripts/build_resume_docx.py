# -*- coding: utf-8 -*-
"""生成美化版 Word 简历（一页版，带照片）"""
import os
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC_PHOTO = r"C:/Users/陈词年/Desktop/简历/_photo_extracted.png"
OUT = r"C:/Users/陈词年/Desktop/简历/陈如忆-简历-初级版.docx"

THEME = RGBColor(0x2C, 0x3E, 0x50)
THEME_LIGHT = "E8EEF4"
TEXT_DARK = RGBColor(0x33, 0x33, 0x33)
TEXT_GRAY = RGBColor(0x66, 0x66, 0x66)


def set_font(run, name="宋体", size=10.5, bold=False, color=None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=2, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_left_border(paragraph, color_hex="2C3E50", size="28", space="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), space)
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def remove_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "nil")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_run(p, text, size=10.5, bold=False, color=None, name="宋体"):
    run = p.add_run(text)
    set_font(run, name=name, size=size, bold=bold, color=color)
    return run


def add_bullet(doc, text, size=10.5, indent=14):
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=2, line=1.25)
    p.paragraph_format.left_indent = Pt(indent)
    add_run(p, text, size=size)
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=6, after=4, line=1.2)
    set_shading(p, THEME_LIGHT)
    set_left_border(p, color_hex="2C3E50", size="28", space="6")
    p.paragraph_format.left_indent = Pt(6)
    add_run(p, "  " + text, size=12, bold=True, color=THEME, name="黑体")
    return p


# ===== 照片处理：圆角 + 蓝边 =====
def stylize_photo(src, out, target_w=240):
    img = Image.open(src).convert("RGBA")
    ratio = target_w / img.width
    new_size = (target_w, int(img.height * ratio))
    img = img.resize(new_size, Image.LANCZOS)
    radius = 12
    mask = Image.new("L", new_size, 0)
    d = ImageDraw.Draw(mask)
    d.rounded_rectangle((0, 0, new_size[0]-1, new_size[1]-1), radius=radius, fill=255)
    img.putalpha(mask)
    bordered = Image.new("RGBA", (new_size[0]+8, new_size[1]+8), (0, 0, 0, 0))
    bordered.paste(img, (4, 4), img)
    d2 = ImageDraw.Draw(bordered)
    d2.rounded_rectangle((0, 0, bordered.width-1, bordered.height-1),
                         radius=radius+4, outline=(44, 62, 80, 255), width=3)
    bordered.save(out, "PNG")
    return out

PHOTO_STYLED = r"C:/Users/陈词年/Desktop/简历/_photo_styled.png"
stylize_photo(SRC_PHOTO, PHOTO_STYLED, target_w=240)


# ===== 构建 docx =====
doc = Document()
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(12)
section.bottom_margin = Mm(12)
section.left_margin = Mm(16)
section.right_margin = Mm(16)

style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# ===== 顶部：照片 + 姓名信息 =====
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_table_borders(header_table)
header_table.autofit = False
header_table.columns[0].width = Mm(28)
header_table.columns[1].width = Mm(150)
row = header_table.rows[0]
row.cells[0].width = Mm(28)
row.cells[1].width = Mm(150)

photo_cell = row.cells[0]
photo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
photo_para = photo_cell.paragraphs[0]
photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(photo_para, before=0, after=0)
photo_para.add_run().add_picture(PHOTO_STYLED, width=Mm(23))

info_cell = row.cells[1]
info_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
info_cell.paragraphs[0].text = ""

p_name = info_cell.paragraphs[0]
set_spacing(p_name, before=0, after=1, line=1.1)
add_run(p_name, "陈如忆", size=21, bold=True, color=THEME, name="黑体")

p_obj = info_cell.add_paragraph()
set_spacing(p_obj, before=1, after=2, line=1.2)
add_run(p_obj, "求职意向：", size=10.5, bold=True, color=TEXT_DARK)
add_run(p_obj, "软件开发 / 软件测试 / 数据分析（初级岗位）", size=10.5, color=TEXT_DARK)

p_meta = info_cell.add_paragraph()
set_spacing(p_meta, before=0, after=0, line=1.2)
add_run(p_meta, "女  |  23 岁  |  本科  |  共青团员", size=9, color=TEXT_GRAY)

p_contact = info_cell.add_paragraph()
set_spacing(p_contact, before=0, after=0, line=1.2)
add_run(p_contact, "电话：19516126370    邮箱：2401671681@qq.com    GitHub：github.com/KAKAstarrr/f1-light-app", size=9, color=TEXT_GRAY)

# ===== 教育经历 =====
add_section_title(doc, "教育经历")
p_edu = doc.add_paragraph()
set_spacing(p_edu, before=3, after=2, line=1.25)
add_run(p_edu, "2021.09 – 2025.06   徐州工程学院   计算机科学与技术（嵌入式培养）· 本科",
        size=10.5, bold=True, color=TEXT_DARK)
for it in [
    "主修课程：数据结构、操作系统、数据库原理及应用、计算机网络、软件工程基础",
    "校 ACM 社团成员，多次参加团体编程竞赛；院办公室部门成员，参与活动策划与新闻稿撰写",
    "成绩优异，多次获得校特等奖学金",
]:
    add_bullet(doc, "• " + it)

# ===== 项目经历 =====
add_section_title(doc, "项目经历")

p1 = doc.add_paragraph()
set_spacing(p1, before=3, after=1, line=1.2)
add_run(p1, "F1 赛事数据平台（个人项目）", size=10.5, bold=True, color=TEXT_DARK)
add_run(p1, "    2026.07 – 至今", size=9, color=TEXT_GRAY)

p1s = doc.add_paragraph()
set_spacing(p1s, before=0, after=3, line=1.2)
add_run(p1s, "FastAPI + Vue 3 + SQLAlchemy + pandas / numpy + ECharts + XGBoost",
        size=9, color=THEME)
add_run(p1s, "  ·  借助 AI 编程助手辅助开发", size=9, color=TEXT_GRAY)

for it in [
    "开发 39 个 RESTful API，覆盖赛程、成绩、遥测分析、AI 预测、Fantasy 游戏、投票等模块",
    "设计三级缓存，将遥测数据首次加载 60 秒以上降至缓存命中后秒级返回；用 pandas 处理 20Hz 数据并降采样至约 200 点，ECharts 实现 6 图层可视化",
    "用 XGBoost 训练夺冠预测模型（19 特征，2018–2025 年 3400+ 条数据），Top1 命中率 41.7%、Top3 命中率 91.7%，优于规则模型（33.3%）",
    "实现 Fantasy 积分游戏（动态定价、芯片、联盟）与 JWT 用户鉴权，基于 SQLite 8 张表",
]:
    add_bullet(doc, "• " + it)

p2 = doc.add_paragraph()
set_spacing(p2, before=6, after=1, line=1.2)
add_run(p2, "携程旅行 App 定制旅行功能系统分析与设计（课程项目）",
        size=10.5, bold=True, color=TEXT_DARK)
add_run(p2, "    2024.04 – 2024.06", size=9, color=TEXT_GRAY)

p2s = doc.add_paragraph()
set_spacing(p2s, before=0, after=3, line=1.2)
add_run(p2s, "项目负责人  |  Astah（UML 建模工具）", size=9, color=TEXT_GRAY)

for it in [
    "用 UML（业务用例图、业务泳道图、类图、顺序图）完成「定制需求提交」端到端分析与设计",
    "撰写用例规约、产出设计文档，驱动小组接口对接，项目获评优秀",
]:
    add_bullet(doc, "• " + it)

# ===== 技能 =====
add_section_title(doc, "技能")
for label, content in [
    ("编程语言", "熟悉 C / C++、Python，掌握 SQL（MySQL）"),
    ("Web 开发", "使用过 FastAPI、Vue 3、Element Plus、SQLAlchemy"),
    ("数据与测试", "pandas / numpy 数据清洗统计，了解 XGBoost；熟悉测试生命周期与黑盒测试方法"),
    ("工具", "Git、Linux、Excel，会使用 AI 编程助手辅助开发"),
]:
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=2, line=1.3)
    add_run(p, f"{label}：", size=10.5, bold=True, color=THEME)
    add_run(p, content, size=10.5, color=TEXT_DARK)

# ===== 荣誉证书 =====
add_section_title(doc, "荣誉证书")
add_bullet(doc,
    "• 大学英语四级（CET-4）、蓝桥杯江苏省 C / C++ 程序设计 B 组三等奖、"
    "RoboCom 计算机开发者大赛省赛优秀奖、团体程序设计天梯赛省高校三等奖")

# ===== 自我评价 =====
add_section_title(doc, "自我评价")
for label, content in [
    ("基础扎实", "计算机科班，ACM 社团 + 省级竞赛获奖 + 多次特等奖学金，数据结构与算法功底扎实"),
    ("有实践", "完成全栈个人项目与课程项目，能独立定位和解决问题"),
    ("肯学习", "对新技术保持好奇，愿意从基础岗位做起、踏实成长"),
]:
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=2, line=1.3)
    add_run(p, f"{label}：", size=10.5, bold=True, color=THEME)
    add_run(p, content, size=10.5, color=TEXT_DARK)

doc.save(OUT)
print("saved:", OUT)

if os.path.exists(PHOTO_STYLED):
    try:
        os.remove(PHOTO_STYLED)
    except OSError:
        pass
