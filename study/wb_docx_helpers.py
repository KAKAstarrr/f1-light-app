# -*- coding: utf-8 -*-
"""python-docx 样式辅助模块：标题体系 / 代码块底纹 / 问答格式 / 表格 / 目录字段"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "宋体"
HEAD_FONT = "微软雅黑"
CODE_FONT = "Consolas"

# 主题色
C_DARK = (0x1F, 0x3B, 0x63)   # 深蓝
C_MID = (0x2E, 0x59, 0x8C)    # 中蓝
C_ACC = (0x3E, 0x6B, 0xA8)    # 亮蓝
C_ORANGE = (0xB0, 0x4A, 0x2E) # 问题橙
C_GREEN = (0x2E, 0x6B, 0x4A)  # 回答绿
C_CODE = (0x33, 0x33, 0x33)
C_GRAY = (0x66, 0x66, 0x66)


def set_run_font(run, name=BODY_FONT, size=10.5, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def shade_paragraph(paragraph, fill="F5F5F5"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, 17, True, C_DARK)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '10')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1F3B63')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, 13.5, True, C_MID)
    return p


def add_h3(doc, text, color=C_ACC):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, HEAD_FONT, 11.5, True, color)
    return p


def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.3
    if indent:
        p.paragraph_format.first_line_indent = Pt(21)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, 10.5)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, BODY_FONT, 10.5, True)
    run = p.add_run(text)
    set_run_font(run, BODY_FONT, 10.5)
    return p


def add_code_block(doc, code, title=None):
    if title:
        pt = doc.add_paragraph()
        pt.paragraph_format.space_before = Pt(4)
        pt.paragraph_format.space_after = Pt(2)
        rt = pt.add_run(title)
        set_run_font(rt, HEAD_FONT, 10, True, C_ORANGE)
    lines = code.strip('\n').split('\n')
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Pt(12)
        shade_paragraph(p, 'F2F2F2')
        run = p.add_run(line if line else ' ')
        set_run_font(run, CODE_FONT, 9, color=C_CODE)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    sp.paragraph_format.line_spacing = 0.5
    set_run_font(sp.add_run(''), BODY_FONT, 3)


def add_qa(doc, q, a):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.25
    rq = p.add_run("Q：")
    set_run_font(rq, HEAD_FONT, 10.5, True, C_ORANGE)
    rq2 = p.add_run(q)
    set_run_font(rq2, BODY_FONT, 10.5, True)
    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(5)
    ap.paragraph_format.line_spacing = 1.25
    ra = ap.add_run("A：")
    set_run_font(ra, HEAD_FONT, 10.5, True, C_GREEN)
    ra2 = ap.add_run(a)
    set_run_font(ra2, BODY_FONT, 10.5)
    return


def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, HEAD_FONT, 10, True, (0xFF, 0xFF, 0xFF))
        # 表头底色
        tcPr = hdr[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '2E598C')
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            set_run_font(r, BODY_FONT, 9.5)
            p.paragraph_format.space_after = Pt(1)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    # 表后空段
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)
    set_run_font(sp.add_run(''), BODY_FONT, 4)
    return t


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def render_unit(doc, unit, with_h1=False):
    """统一渲染一个知识单元。unit 字段：
    id / title / concept(列表，支持 ('body'|'bullet'|'code'|'h3', ...)) /
    code(dict 标题->代码) / pits(列表，支持 str 或 (前缀, 文本)) / qa(列表 (q,a)) / table(headers, rows, widths)"""
    if with_h1:
        add_h1(doc, unit['title'])
    else:
        add_h2(doc, f"{unit['id']}  {unit['title']}")

    if unit.get('concept'):
        add_h3(doc, '知识点详解')
        for item in unit['concept']:
            if isinstance(item, str):
                add_body(doc, item)
            elif item[0] == 'body':
                add_body(doc, item[1])
            elif item[0] == 'bullet':
                bp = item[2] if len(item) > 2 else None
                add_bullet(doc, item[1], bold_prefix=bp)
            elif item[0] == 'code':
                add_code_block(doc, item[1], item[2] if len(item) > 2 else None)
            elif item[0] == 'h3':
                add_h3(doc, item[1])
            elif item[0] == 'table':
                add_table(doc, item[1], item[2], item[3] if len(item) > 3 else None)

    if unit.get('table'):
        add_h3(doc, '要点速览')
        add_table(doc, unit['table'][0], unit['table'][1], unit['table'][2] if len(unit['table']) > 2 else None)

    if unit.get('code'):
        add_h3(doc, '项目实战代码')
        for title, code in unit['code'].items():
            add_code_block(doc, code, title)

    if unit.get('pits'):
        add_h3(doc, '踩坑记录', (0xA8, 0x3A, 0x2E))
        for pit in unit['pits']:
            if isinstance(pit, tuple):
                add_bullet(doc, pit[1], bold_prefix='【' + pit[0] + '】 ')
            else:
                add_bullet(doc, pit)

    if unit.get('qa'):
        add_h3(doc, '面试问答', (0x2E, 0x6B, 0x4A))
        for q, a in unit['qa']:
            add_qa(doc, q, a)
    return
