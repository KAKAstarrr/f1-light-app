# -*- coding: utf-8 -*-
"""F1 项目知识点整合手册生成器
用法：python build_knowledge_word.py
产出：study/F1项目知识点整合手册.docx
"""
import os
import sys
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from wb_docx_helpers import (
    set_run_font, add_h1, add_h2, add_h3, add_body, add_bullet,
    add_code_block, add_qa, add_table, add_page_break, render_unit,
    HEAD_FONT, BODY_FONT,
)
from wb_content_part0_1 import UNITS as U01, PART0_INTRO
from wb_content_frontend import UNITS as UFE, FRONTEND_INTRO
from wb_content_part2 import UNITS as U02, PART2_INTRO
from wb_content_part3 import UNITS as U03, PART3_INTRO
from wb_content_part4 import UNITS as U04, PART4_INTRO
from wb_content_ds import UNITS as UDS, DS_INTRO
from wb_content_interview import ROLES, COMMON_QA, INTERVIEW_INTRO

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "F1项目知识点整合手册.docx")


def build_cover(doc):
    for _ in range(4):
        p = doc.add_paragraph()
        set_run_font(p.add_run(""), BODY_FONT, 14)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("F1 数据交互应用")
    set_run_font(r, HEAD_FONT, 30, True, (0x1F, 0x3B, 0x63))
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("全栈 + AI 知识点整合手册")
    set_run_font(r2, HEAD_FONT, 24, True, (0x2E, 0x59, 0x8C))
    for _ in range(2):
        p = doc.add_paragraph()
        set_run_font(p.add_run(""), BODY_FONT, 10)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("覆盖阶段 0-4 + 扩展模块 + 数据科学 + Streamlit\n"
                     "含多岗位面试题库（数据分析 / 产品经理 / 后端 / 前端 / 全栈 / AI-ML）")
    set_run_font(rs, HEAD_FONT, 13, False, (0x66, 0x66, 0x66))
    for _ in range(8):
        p = doc.add_paragraph()
        set_run_font(p.add_run(""), BODY_FONT, 10)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run(f"项目：f1_light_app\n"
                      f"技术栈：FastAPI · Vue3 · XGBoost · FastF1 · Ergast · SQLite\n"
                      f"整理日期：{datetime.now().strftime('%Y-%m-%d')}")
    set_run_font(rm, BODY_FONT, 11, False, (0x99, 0x99, 0x99))
    add_page_break(doc)


def build_toc_overview(doc):
    add_h1(doc, "文档导览")
    add_body(doc, "本手册由 study/ 目录下 9 份分阶段知识点文档整合而成，采用统一的"
                  "「知识点详解 + 项目实战代码 + 踩坑记录 + 面试问答」四段式结构，"
                  "覆盖从环境搭建到 AI 模型上线的完整技术链路。")
    add_table(doc, ["章节", "内容", "对应岗位"], [
        ["第一部分", "阶段 0：环境与基础（FastAPI/Ergast/FastF1/Conda/Git/Pydantic）", "后端、全栈"],
        ["第二部分", "阶段 1：后端基础数据（三级缓存/分层/序列化/轮胎策略/兜底）", "后端、全栈"],
        ["第三部分", "前端 Vue3 全面知识点（组合式 API/Pinia/Router/Composables）", "前端、全栈"],
        ["第四部分", "阶段 2：遥测分析 + 扩展（B1-B6：对比/分布/赛道地图/天气）", "数据分析、前端"],
        ["第五部分", "阶段 3：数据库/认证/Fantasy/规则预测", "后端、产品经理"],
        ["第六部分", "阶段 4：XGBoost AI 预测（采集/特征/训练/推理/SHAP）", "AI/ML、数据分析"],
        ["第七部分", "数据科学基础（NumPy/Pandas/Matplotlib）", "数据分析"],
        ["第八部分", "Streamlit 快速原型开发", "数据分析、全栈"],
        ["第九部分", "多岗位面试题库（6 类岗位 + 通用题）", "全部岗位"],
    ], [2.6, 9.4, 3.0])
    add_body(doc, "使用建议：按目标岗位跳读——后端岗重点看第一、二、五部分；前端岗看第三部分；"
                  "数据岗看第四、七、八部分；AI 岗看第六部分；产品岗看第九部分产品经理题库；"
                  "面试前通读第九部分并按 STAR 法则用自己的话复述。")
    add_page_break(doc)


def build_part(doc, part_no, title, units, intro):
    add_h1(doc, f"{part_no}  {title}")
    add_body(doc, intro, indent=False)
    for u in units:
        render_unit(doc, u)


def build_interview(doc):
    add_h1(doc, "第九部分  多岗位面试题库")
    add_body(doc, INTERVIEW_INTRO, indent=False)
    for role in ROLES:
        add_h2(doc, f"8.{ROLES.index(role) + 1}  {role['role']} 岗位")
        add_h3(doc, "面试官关注点")
        for f in role['focus']:
            add_bullet(doc, f)
        add_h3(doc, "核心面试题（含答题要点）")
        for i, (q, pts) in enumerate(role['questions'], 1):
            add_qa(doc, f"{role['role']} · 第 {i} 题：{q}", pts[0])
            for extra in pts[1:]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run("    ▸ " + extra)
                set_run_font(r, BODY_FONT, 10, False, (0x55, 0x55, 0x55))
        add_h3(doc, "简历亮点建议")
        for r_line in role['resume']:
            add_bullet(doc, r_line)
    add_h2(doc, "8.7  通用高频题（所有岗位）")
    for q, a in COMMON_QA:
        add_qa(doc, q, a)


def main():
    doc = Document()
    # 页边距
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    build_cover(doc)
    build_toc_overview(doc)

    # 第一、二部分（阶段 0 + 阶段 1）
    build_part(doc, "第一部分", "阶段 0：环境与基础", U01[:7],
               "本部分覆盖 FastAPI 路由与参数、CORS 中间件、Ergast API 数据解包、"
               "FastF1 对象模型、Conda 环境、Git 版本控制、Pydantic 数据校验。")
    build_part(doc, "第二部分", "阶段 1：后端基础数据", U01[7:],
               "本部分覆盖三层架构与 API 分层、三级缓存架构、数据序列化（timedelta/numpy）、"
               "最快圈提取（groupby+idxmin）、轮胎策略重建（Stint 分组）、异常兜底（None 模式）、"
               "缓存键设计。")
    build_part(doc, "第三部分", "前端：Vue3 全面知识点", UFE,
               FRONTEND_INTRO)
    build_part(doc, "第四部分", "阶段 2：遥测分析 + 扩展（B1-B6）", U02,
               PART2_INTRO)
    build_part(doc, "第五部分", "阶段 3：数据库 / 认证 / Fantasy / 规则预测", U03,
               PART3_INTRO)
    build_part(doc, "第六部分", "阶段 4：XGBoost AI 预测", U04,
               PART4_INTRO)
    build_part(doc, "第七部分", "数据科学基础（NumPy / Pandas / Matplotlib）", UDS[:4],
               "本部分为数据科学三件套知识点，面向数据分析师与全栈岗位。")
    build_part(doc, "第八部分", "Streamlit 快速原型开发", UDS[4:],
               "本部分为 Streamlit 快速原型开发知识点：「纯 Python 出 Web 看板」的 MVP 实践。")
    build_interview(doc)

    doc.save(OUT_PATH)
    print(f"[OK] 已生成：{OUT_PATH}")
    print(f"[OK] 文件大小：{os.path.getsize(OUT_PATH) / 1024:.1f} KB")


if __name__ == "__main__":
    main()
