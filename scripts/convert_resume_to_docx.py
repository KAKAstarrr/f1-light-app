# -*- coding: utf-8 -*-
"""将 PDF 简历按页渲染为高清图片，嵌入 Word，保留原版式与照片"""
from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

PDF = r"C:/Users/陈词年/Desktop/简历/陈如忆-运营.pdf"
OUT = r"C:/Users/陈词年/Desktop/简历/陈如忆-运营.docx"
TMP_DIR = r"C:/Users/陈词年/Desktop/简历"

# 1. 渲染 PDF 每页到高清 PNG（保留照片和所有装饰）
import pypdfium2 as pdfium
pdf = pdfium.PdfDocument(PDF)
pages_png = []
for i in range(len(pdf)):
    page = pdf[i]
    bitmap = page.render(scale=3.0)  # 高清，3x
    img = bitmap.to_pil()
    out_png = os.path.join(TMP_DIR, f"_pdf_page_{i+1}.png")
    img.save(out_png, optimize=True)
    pages_png.append(out_png)
    print(f"rendered page {i+1} -> {out_png}  size={img.size}")

# 2. 创建 Word：A4，每页一张图，铺满页面（带少量页边距）
doc = Document()

# 设为 A4，页边距设小一点，让图片尽量接近原 PDF 比例
section = doc.sections[0]
section.page_width = Mm(210)
section.page_height = Mm(297)
section.top_margin = Mm(8)
section.bottom_margin = Mm(8)
section.left_margin = Mm(10)
section.right_margin = Mm(10)

# 默认字体设置（图片页不需要太多文字样式）
style = doc.styles["Normal"]
style.font.size = Pt(10.5)

for png_path in pages_png:
    # 每页一张图，居中，最大宽度接近页面宽度
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    # 让图片宽度 = 内容区宽度（A4 宽 210 - 左右边距 ≈ 190mm）
    run.add_picture(png_path, width=Mm(190))

    # 最后一页不加 page break
    if png_path is not pages_png[-1]:
        from docx.enum.text import WD_BREAK
        p.add_run().add_break(WD_BREAK.PAGE)

doc.save(OUT)
print("saved:", OUT)

# 3. 清理临时预览图（保留最终 _pdf_page_*.png 直到 docx 生成成功）
for p in pages_png:
    if os.path.exists(p):
        os.remove(p)
        print("cleaned", p)